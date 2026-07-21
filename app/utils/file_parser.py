"""
File parser utilities: extract text and tables from PDF, DOCX, image, and plain-text files
into a normalised ParsedDocument structure.
"""
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn as docx_qn
from PIL import Image
import fitz
import pytesseract
import logging
logger = logging.getLogger(__name__)

from app.utils.text_normalizer import normalize_text


@dataclass
class ParsedDocument:
    pages: list[tuple[int, str]]
    full_text: str


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

# Return True if two (x0, y0, x1, y1) rectangles overlap within an optional margin.
def _rects_overlap(r1: tuple, r2: tuple, margin: float = 2.0) -> bool:
    return not (
        r1[2] + margin < r2[0]
        or r2[2] + margin < r1[0]
        or r1[3] + margin < r2[1]
        or r2[3] + margin < r1[1]
    )


# Merge adjacent column pairs caused by PDF merged-cell artifacts (handles spanned-cell layouts).
def _merge_complementary_cols(rows: list[list[str]]) -> list[list[str]]:
    if not rows or len(rows[0]) < 2:
        return rows

    n_rows = len(rows)
    n_cols = len(rows[0])
    col_groups: list[list[str]] = []

    col = 0
    while col < n_cols:
        if col + 1 < n_cols:
            # Mergeable = never have two *different* non-empty values in same row
            mergeable = all(
                not (rows[r][col] and rows[r][col + 1]
                     and rows[r][col] != rows[r][col + 1])
                for r in range(n_rows)
            )
            if mergeable:
                merged = [rows[r][col] or rows[r][col + 1] for r in range(n_rows)]
                col_groups.append(merged)
                col += 2
                continue
        col_groups.append([rows[r][col] for r in range(n_rows)])
        col += 1

    n_out = len(col_groups)
    return [[col_groups[c][r] for c in range(n_out)] for r in range(n_rows)]


# Convert a PyMuPDF Table object to a Markdown table string, handling merged-cell artifacts.
def _table_to_markdown(table) -> str:
    try:
        rows = table.extract()
        if not rows:
            return ""

        # Clean and pad all rows to same column count
        ncols = max(len(r) for r in rows)
        cleaned: list[list[str]] = []
        for row in rows:
            padded = (list(row) + [None] * ncols)[:ncols]
            cleaned.append([str(cell or "").strip().replace("\n", " ") for cell in padded])

        if not cleaned:
            return ""

        # Step 1: drop columns that are entirely empty
        non_empty_cols = [
            col for col in range(ncols)
            if any(cleaned[r][col] for r in range(len(cleaned)))
        ]
        if not non_empty_cols:
            return ""
        filtered = [[row[col] for col in non_empty_cols] for row in cleaned]

        # Step 2: merge adjacent complementary column pairs (merged-cell artifact)
        filtered = _merge_complementary_cols(filtered)

        n = len(filtered[0])
        lines = [
            "| " + " | ".join(filtered[0]) + " |",
            "| " + " | ".join("---" for _ in range(n)) + " |",
        ]
        for row in filtered[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except Exception as e:
        logger.warning("_table_to_markdown failed: %s", e)
        return ""


# Wrap non-empty text with **bold** / *italic* markers, preserving surrounding whitespace.
def _apply_markdown_fmt(text: str, is_bold: bool, is_italic: bool) -> str:
    stripped = text.strip()
    if not stripped:
        return text
    prefix = text[: len(text) - len(text.lstrip())]
    suffix = text[len(text.rstrip()) :]
    if is_bold and is_italic:
        return f"{prefix}***{stripped}***{suffix}"
    if is_bold:
        return f"{prefix}**{stripped}**{suffix}"
    if is_italic:
        return f"{prefix}*{stripped}*{suffix}"
    return text


_BOLD_FONT_HINTS = frozenset({"bold", "black", "heavy", "demi", "semibold", "extrabold"})


# Convert a PyMuPDF dict-mode text block to Markdown, detecting bold/italic from span flags and font names.
def _block_to_formatted_text(block: dict) -> str:
    line_texts: list[str] = []
    for line in block.get("lines", []):
        span_parts: list[str] = []
        for span in line.get("spans", []):
            raw = span.get("text", "")
            if not raw:
                continue
            flags: int = span.get("flags", 0)
            font: str = (span.get("font") or "").lower()
            is_bold = bool(flags & 16) or any(h in font for h in _BOLD_FONT_HINTS)
            is_italic = bool(flags & 2) or "italic" in font or "oblique" in font
            span_parts.append(_apply_markdown_fmt(raw, is_bold, is_italic))
        line_text = "".join(span_parts).rstrip("\n")
        if line_text.strip():
            line_texts.append(line_text)
    return "\n".join(line_texts).strip()


# Remove OCR noise: keep only lines that have ≥2 meaningful words (each ≥2 alpha chars).
# Returns cleaned word-only text (strips pipe chars and other OCR artifacts from each line).
def _clean_ocr_text(raw: str) -> str:
    good: list[str] = []
    for line in raw.splitlines():
        line = line.strip()
        if not line:
            continue
        valid_words: list[str] = []
        for token in line.split():
            # Strip leading/trailing non-alphanumeric characters
            core = re.sub(r'^[^\w]+|[^\w]+$', '', token)
            if len(core) < 2:
                continue
            alpha = sum(1 for c in core if c.isalpha())
            # Need ≥2 alphabetic chars and ≥50% of the token to be alphabetic
            if alpha >= 2 and alpha / len(core) >= 0.5:
                valid_words.append(core)
        if len(valid_words) >= 2:
            # Return cleaned words only (strips | and other non-word artifacts)
            good.append(' '.join(valid_words))
    return '\n'.join(good)


# If OCR lines look like flowchart labels (short, uniform), join with → to indicate sequence.
def _maybe_join_as_flow(lines: list[str]) -> str:
    if len(lines) < 3:
        return '\n'.join(lines)
    word_counts = [len(l.split()) for l in lines]
    avg_words = sum(word_counts) / len(word_counts)
    max_words = max(word_counts)
    # Flowchart heuristic: all labels short (avg ≤ 5 words, no outlier > 2× average)
    if avg_words <= 5 and max_words <= max(avg_words * 2, 6):
        return ' → '.join(lines)
    return '\n'.join(lines)


# OCR a rectangular region of a PDF page.
# Runs two PSM passes (auto + sparse) at high zoom, deduplicates, returns cleaned text.
def _ocr_page_region(page: fitz.Page, bbox: tuple, zoom: float = 4.0) -> str:
    try:
        from PIL import ImageEnhance, ImageFilter
        clip = fitz.Rect(bbox)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, clip=clip, colorspace=fitz.csGRAY)
        img = Image.open(BytesIO(pix.tobytes("png")))
        # Enhance contrast + sharpen to help OCR on faint/compressed text
        img = ImageEnhance.Contrast(img).enhance(1.8)
        img = img.filter(ImageFilter.SHARPEN)

        seen: set[str] = set()
        collected: list[str] = []
        # PSM 3 = full auto (good for structured diagrams)
        # PSM 11 = sparse text (good for scattered labels)
        for psm in (3, 11):
            raw = pytesseract.image_to_string(img, lang="vie+eng", config=f"--oem 1 --psm {psm}")
            for line in _clean_ocr_text(raw).splitlines():
                key = re.sub(r'\s+', ' ', line.strip().lower())
                if key and key not in seen:
                    seen.add(key)
                    collected.append(line.strip())
        # Join short uniform labels with → (flowchart heuristic) to signal sequential flow to LLM
        return _maybe_join_as_flow(collected)
    except Exception as e:
        logger.debug("_ocr_page_region failed: %s", e)
        return ""


# Extract text from one PDF page: tables → Markdown, other text preserves bold/italic; no duplication of table regions.
# Image/drawing blocks (type==1) are OCR-ed so flowcharts and diagrams are captured.
def _extract_pdf_page(page: fitz.Page) -> str:
    try:
        tab_finder = page.find_tables()
        tables = tab_finder.tables if tab_finder else []
    except Exception:
        tables = []

    logger.info("PDF page %s: find_tables detected %d table(s)", page.number + 1, len(tables))
    table_rects = [tuple(t.bbox) for t in tables]

    # Use dict mode to get per-span font/flag information
    page_dict = page.get_text("dict")
    items: list[tuple[float, str]] = []

    for block in page_dict.get("blocks", []):
        btype = block.get("type", 0)
        bx0, by0, bx1, by1 = block["bbox"]

        if btype == 1:
            # Image/drawing block — OCR it to capture flowchart / diagram text.
            # Skip if the region is tiny (icons, bullets).
            width, height = bx1 - bx0, by1 - by0
            if width < 15 or height < 10:
                continue
            if table_rects and any(_rects_overlap((bx0, by0, bx1, by1), tr) for tr in table_rects):
                continue
            ocr_text = _ocr_page_region(page, (bx0, by0, bx1, by1))
            if ocr_text:
                items.append((by0, ocr_text))
            continue

        # type == 0: regular text block
        if table_rects and any(_rects_overlap((bx0, by0, bx1, by1), tr) for tr in table_rects):
            continue                    # covered by table markdown below
        block_text = _block_to_formatted_text(block)
        if block_text:
            items.append((by0, block_text))

    # Add each table as Markdown at its vertical position
    for table in tables:
        md = _table_to_markdown(table)
        if md:
            items.append((table.bbox[1], md))

    # Sort by vertical position to preserve reading order
    items.sort(key=lambda x: x[0])
    return "\n\n".join(text for _, text in items)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

# Extract text from a w:drawing element (SmartArt, shapes, text boxes) by collecting
# all a:t (DrawingML text run) nodes recursively. Returns empty string if nothing found.
def _drawing_to_text(drawing_elem) -> str:
    # DrawingML text nodes live under the "a" namespace
    _A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    texts: list[str] = []
    for node in drawing_elem.iter(f"{{{_A_NS}}}t"):
        t = (node.text or "").strip()
        if t:
            texts.append(t)
    return " → ".join(texts) if texts else ""


# Parse a DOCX file to text preserving paragraphs, tables (as Markdown), and drawings/SmartArt in document order.
def _docx_to_text(raw_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(raw_bytes))

    parts: list[str] = []
    for child in doc.element.body:
        tag = child.tag

        if tag == docx_qn("w:p"):
            # Top-level paragraph: collect runs with bold/italic detection
            run_parts: list[str] = []
            for run in child.iter(docx_qn("w:r")):
                rpr = run.find(docx_qn("w:rPr"))
                is_bold = rpr is not None and rpr.find(docx_qn("w:b")) is not None
                is_italic = rpr is not None and rpr.find(docx_qn("w:i")) is not None
                raw = "".join(node.text or "" for node in run.iter(docx_qn("w:t")))
                if raw:
                    run_parts.append(_apply_markdown_fmt(raw, is_bold, is_italic))
            text = "".join(run_parts).strip()
            if text:
                parts.append(text)
            # Also capture any drawings embedded inside this paragraph
            for drawing in child.iter(docx_qn("w:drawing")):
                d_text = _drawing_to_text(drawing)
                if d_text:
                    parts.append(d_text)

        elif tag == docx_qn("w:tbl"):
            # Table: convert directly from XML to avoid nested-para confusion
            rows: list[list[str]] = []
            for row_elem in child.iter(docx_qn("w:tr")):
                cells: list[str] = []
                for cell_elem in row_elem.findall(".//" + docx_qn("w:tc")):
                    cell_text = "".join(
                        node.text or "" for node in cell_elem.iter(docx_qn("w:t"))
                    )
                    cells.append(cell_text.strip().replace("\n", " "))
                if cells:
                    rows.append(cells)

            if rows:
                ncols = max(len(r) for r in rows)
                header = (rows[0] + [""] * ncols)[:ncols]
                lines = [
                    "| " + " | ".join(header) + " |",
                    "| " + " | ".join("---" for _ in range(ncols)) + " |",
                ]
                for row in rows[1:]:
                    padded = (row + [""] * ncols)[:ncols]
                    lines.append("| " + " | ".join(padded) + " |")
                parts.append("\n".join(lines))

        elif tag == docx_qn("w:drawing"):
            # Top-level drawing (rare but possible)
            d_text = _drawing_to_text(child)
            if d_text:
                parts.append(d_text)

    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

# Dispatch raw file bytes to the appropriate parser and return a normalised ParsedDocument.
def parse_file_bytes(raw_bytes: bytes, filename: str, mime_type: str) -> ParsedDocument:
    ext = Path(filename).suffix.lower()
    pages: list[tuple[int, str]] = []

    if ext in {".txt", ".md", ".csv", ".log"} or mime_type.startswith("text/"):
        text = raw_bytes.decode("utf-8", errors="ignore")
        pages = [(1, text)]

    elif ext == ".pdf" or mime_type == "application/pdf":
        doc = fitz.open(stream=BytesIO(raw_bytes), filetype="pdf")
        for i, page in enumerate(doc, start=1):
            pages.append((i, _extract_pdf_page(page)))
        doc.close()

    elif ext == ".docx" or mime_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }:
        text = _docx_to_text(raw_bytes)
        pages = [(1, text)]

    elif ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp"} or mime_type.startswith("image/"):
        img = Image.open(BytesIO(raw_bytes))
        text = pytesseract.image_to_string(img)
        pages = [(1, text)]

    else:
        text = raw_bytes.decode("utf-8", errors="ignore")
        pages = [(1, text)]

    pages = [(p, normalize_text(t)) for p, t in pages if t and normalize_text(t)]
    full_text = normalize_text("\n\n".join(t for _, t in pages))
    if not full_text:
        raise ValueError("No extractable text in file")
    return ParsedDocument(pages=pages, full_text=full_text)


# Strip repeated header/footer lines that appear on at least 'threshold' pages.
def _strip_headers_footers(pages: list[tuple[int, str]], threshold: int = 3) -> list[tuple[int, str]]:
    if len(pages) < 2:
        return pages

    from collections import Counter

    N = 3
    top_lines: list[list[str]] = []
    bot_lines: list[list[str]] = []
    for _, text in pages:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        top_lines.append(lines[:N])
        bot_lines.append(lines[-N:])

    top_counter: Counter = Counter()
    for lines in top_lines:
        for line in lines:
            if len(line) > 5:
                top_counter[line] += 1

    bot_counter: Counter = Counter()
    for lines in bot_lines:
        for line in lines:
            if len(line) > 5:
                bot_counter[line] += 1

    noise_lines = (
        {line for line, count in top_counter.items() if count >= threshold}
        | {line for line, count in bot_counter.items() if count >= threshold}
    )
    if not noise_lines:
        return pages

    logger.info("Detected %d header/footer patterns to strip", len(noise_lines))
    return [
        (page_no, "\n".join(l for l in text.splitlines() if l.strip() not in noise_lines).strip())
        for page_no, text in pages
    ]
